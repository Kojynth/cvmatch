"""
CV Parsers - Parsers universels avec support OCR offline
=======================================================

Parsers robustes pour tous formats (PDF, DOCX, ODT, images) avec
OCR local intégré et preprocessing intelligent des documents scannés.
"""

import os
import io
import tempfile
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass
from pathlib import Path
from loguru import logger

# PDF parsing
try:
    import fitz  # PyMuPDF (prioritaire)
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber  # Fallback PDF
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# DOCX/ODT parsing
try:
    from docx import Document
    from docx.document import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from odf.opendocument import load as odf_load
    from odf.text import P
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False

# OCR et preprocessing images
try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import cv2
    import numpy as np
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False

# Utilitaires
from .cv_analyzer import TextBlock, BoundingBox


@dataclass
class DocumentPage:
    """Page d'un document avec métadonnées."""
    page_number: int
    width: float
    height: float
    text_blocks: List[TextBlock]
    raw_text: str
    image_path: Optional[str] = None  # Pour OCR
    ocr_confidence: Optional[float] = None


@dataclass
class ParsedDocument:
    """Document parsé complet."""
    file_path: str
    file_type: str
    pages: List[DocumentPage]
    total_pages: int
    metadata: Dict[str, Any]
    parsing_method: str
    parsing_confidence: float
    errors: List[str]


class ImagePreprocessor:
    """Préprocesseur d'images pour améliorer l'OCR."""
    
    def __init__(self):
        self.opencv_available = OPENCV_AVAILABLE
    
    def preprocess_image(self, image_path: str, target_dpi: int = 300) -> str:
        """
        Préprocesse une image pour optimiser l'OCR.
        
        Args:
            image_path: Chemin vers l'image
            target_dpi: DPI cible pour l'OCR
            
        Returns:
            Chemin vers l'image préprocessée
        """
        if not self.opencv_available:
            logger.warning("OpenCV non disponible, preprocessing limité")
            return image_path
        
        try:
            # Charger image
            image = cv2.imread(image_path)
            if image is None:
                logger.error(f"Impossible de charger l'image: {image_path}")
                return image_path
            
            logger.info("🖼️ Preprocessing image [FILENAME]")
            
            # 1. Conversion en niveaux de gris
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
            
            # 2. Débruitage
            denoised = cv2.medianBlur(gray, 3)
            
            # 3. Correction d'inclinaison (deskew)
            deskewed = self._deskew_image(denoised)
            
            # 4. Binarisation adaptative
            binary = cv2.adaptiveThreshold(
                deskewed, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # 5. Dilatation légère pour améliorer lisibilité
            kernel = np.ones((1,1), np.uint8)
            processed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
            # 6. Redimensionnement si nécessaire
            height, width = processed.shape
            current_dpi = self._estimate_dpi(width, height)
            
            if current_dpi < target_dpi:
                scale_factor = target_dpi / current_dpi
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                processed = cv2.resize(processed, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
            
            # 7. Sauvegarder image préprocessée
            temp_dir = tempfile.mkdtemp()
            preprocessed_path = os.path.join(temp_dir, f"preprocessed_{Path(image_path).name}")
            cv2.imwrite(preprocessed_path, processed)
            
            logger.info(f"✅ Image préprocessée: {preprocessed_path}")
            return preprocessed_path
            
        except Exception as e:
            logger.error(f"Erreur preprocessing image: {e}")
            return image_path
    
    def _deskew_image(self, image: np.ndarray) -> np.ndarray:
        """Corrige l'inclinaison d'une image."""
        try:
            # Détecter lignes avec transformée de Hough
            edges = cv2.Canny(image, 50, 150, apertureSize=3)
            lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=100)
            
            if lines is not None and len(lines) > 0:
                # Calculer angle moyen des lignes horizontales
                angles = []
                for rho, theta in lines[:10]:  # Prendre les 10 premières lignes
                    angle = theta - np.pi/2
                    angles.append(angle)
                
                # Angle de correction
                skew_angle = np.median(angles) * 180 / np.pi
                
                # Rotation si angle significatif
                if abs(skew_angle) > 0.5:
                    h, w = image.shape
                    center = (w // 2, h // 2)
                    rotation_matrix = cv2.getRotationMatrix2D(center, skew_angle, 1.0)
                    deskewed = cv2.warpAffine(image, rotation_matrix, (w, h), 
                                            flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                    return deskewed
            
            return image
            
        except Exception as e:
            logger.warning(f"Erreur deskew: {e}")
            return image
    
    def _estimate_dpi(self, width: int, height: int) -> float:
        """Estime le DPI d'une image."""
        # Heuristique basique : considérer une page A4 standard
        a4_width_inches = 8.27
        a4_height_inches = 11.69
        
        dpi_x = width / a4_width_inches
        dpi_y = height / a4_height_inches
        
        return (dpi_x + dpi_y) / 2


class OCREngine:
    """Moteur OCR avec Tesseract."""
    
    def __init__(self, language: str = 'fra'):
        self.language = language
        self.tesseract_available = TESSERACT_AVAILABLE
        self.preprocessor = ImagePreprocessor()
        
        # Mapping langues Tesseract
        self.lang_mapping = {
            'fr': 'fra',
            'en': 'eng', 
            'es': 'spa',
            'de': 'deu',
            'it': 'ita',
            'pt': 'por',
            'zh': 'chi_sim',
            'ar': 'ara',
            'he': 'heb'
        }
    
    def extract_text_with_coordinates(self, image_path: str) -> Tuple[List[TextBlock], float]:
        """
        Extrait le texte avec coordonnées via OCR.
        
        Returns:
            Tuple[List[TextBlock], confidence_score]
        """
        if not self.tesseract_available:
            logger.error("Tesseract non disponible pour OCR")
            return [], 0.0
        
        try:
            logger.info("🔍 OCR sur [FILENAME]")
            
            # Preprocessing de l'image
            preprocessed_path = self.preprocessor.preprocess_image(image_path)
            
            # Configuration Tesseract
            lang_code = self.lang_mapping.get(self.language, 'fra')
            custom_config = f'--oem 3 --psm 6 -l {lang_code}'
            
            # OCR avec données détaillées
            image = Image.open(preprocessed_path)
            ocr_data = pytesseract.image_to_data(
                image, 
                config=custom_config,
                output_type=pytesseract.Output.DICT
            )
            
            # Conversion en TextBlocks
            text_blocks = []
            block_id = 0
            
            for i in range(len(ocr_data['text'])):
                text = ocr_data['text'][i].strip()
                confidence = int(ocr_data['conf'][i])
                
                if text and confidence > 30:  # Seuil confiance minimum
                    left = ocr_data['left'][i]
                    top = ocr_data['top'][i]
                    width = ocr_data['width'][i]
                    height = ocr_data['height'][i]
                    
                    bbox = BoundingBox(left, top, left + width, top + height)
                    
                    text_block = TextBlock(
                        id=f"ocr_block_{block_id}",
                        text=text,
                        bbox=bbox,
                        page=1,  # Assumé une page pour image
                        confidence=confidence / 100.0
                    )
                    
                    text_blocks.append(text_block)
                    block_id += 1
            
            # Score de confiance global
            confidences = [int(c) for c in ocr_data['conf'] if int(c) > 0]
            global_confidence = np.mean(confidences) / 100.0 if confidences else 0.0
            
            logger.info(f"✅ OCR terminé: {len(text_blocks)} blocs, confiance: {global_confidence:.2f}")
            
            # Nettoyage fichier temporaire
            if preprocessed_path != image_path and os.path.exists(preprocessed_path):
                try:
                    os.remove(preprocessed_path)
                    # Nettoyer répertoire temporaire si vide
                    temp_dir = os.path.dirname(preprocessed_path)
                    if not os.listdir(temp_dir):
                        os.rmdir(temp_dir)
                except:
                    pass  # Pas critique
            
            return text_blocks, global_confidence
            
        except Exception as e:
            logger.error(f"Erreur OCR: {e}")
            return [], 0.0


class PDFParser:
    """Parser PDF avec PyMuPDF prioritaire et pdfplumber fallback."""
    
    def __init__(self):
        self.pymupdf_available = PYMUPDF_AVAILABLE
        self.pdfplumber_available = PDFPLUMBER_AVAILABLE
    
    def parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse un fichier PDF."""
        if self.pymupdf_available:
            return self._parse_with_pymupdf(file_path)
        elif self.pdfplumber_available:
            return self._parse_with_pdfplumber(file_path)
        else:
            logger.error("Aucun parser PDF disponible (PyMuPDF/pdfplumber)")
            return self._empty_document(file_path, "PDF")
    
    def _parse_with_pymupdf(self, file_path: str) -> ParsedDocument:
        """Parse avec PyMuPDF (prioritaire)."""
        try:
            logger.info("📄 Parse PDF avec PyMuPDF: [FILENAME]")
            
            doc = fitz.open(file_path)
            pages = []
            errors = []
            
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    
                    # Extraire texte avec coordonnées
                    text_blocks = []
                    blocks = page.get_text("dict")["blocks"]
                    
                    block_id = 0
                    for block in blocks:
                        if "lines" in block:  # Bloc texte
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span.get("text", "").strip()
                                    if text:
                                        bbox_list = span["bbox"]
                                        bbox = BoundingBox(
                                            bbox_list[0], bbox_list[1], 
                                            bbox_list[2], bbox_list[3]
                                        )
                                        
                                        text_block = TextBlock(
                                            id=f"pdf_block_{page_num}_{block_id}",
                                            text=text,
                                            bbox=bbox,
                                            page=page_num + 1,
                                            font_size=span.get("size", 12.0),
                                            is_bold=bool(span.get("flags", 0) & 2**4),
                                            font_name=span.get("font", "")
                                        )
                                        
                                        text_blocks.append(text_block)
                                        block_id += 1
                    
                    # Texte brut complet
                    raw_text = page.get_text()
                    
                    # Dimensions page
                    rect = page.rect
                    
                    page_data = DocumentPage(
                        page_number=page_num + 1,
                        width=rect.width,
                        height=rect.height,
                        text_blocks=text_blocks,
                        raw_text=raw_text
                    )
                    
                    pages.append(page_data)
                    
                except Exception as e:
                    error_msg = f"Erreur page {page_num + 1}: {e}"
                    logger.error(error_msg)
                    errors.append(error_msg)
            
            doc.close()
            
            # Métadonnées
            metadata = {
                'creator': 'PyMuPDF',
                'total_text_blocks': sum(len(p.text_blocks) for p in pages),
                'has_images': False  # À améliorer
            }
            
            return ParsedDocument(
                file_path=file_path,
                file_type="PDF",
                pages=pages,
                total_pages=len(pages),
                metadata=metadata,
                parsing_method="PyMuPDF",
                parsing_confidence=0.9 if not errors else 0.7,
                errors=errors
            )
            
        except Exception as e:
            logger.error(f"Erreur PyMuPDF: {e}")
            return self._empty_document(file_path, "PDF")
    
    def _parse_with_pdfplumber(self, file_path: str) -> ParsedDocument:
        """Parse avec pdfplumber (fallback)."""
        try:
            logger.info("📄 Parse PDF avec pdfplumber: [FILENAME]")
            
            import pdfplumber
            
            pages = []
            errors = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Extraire caractères avec coordonnées
                        text_blocks = []
                        chars = page.chars
                        
                        # Grouper caractères en mots/blocs
                        current_text = ""
                        current_bbox = None
                        block_id = 0
                        
                        for char in chars:
                            if char['text'].strip():
                                if current_text and (
                                    not current_bbox or 
                                    abs(char['x0'] - current_bbox[2]) > 5 or  # Espace horizontal
                                    abs(char['y0'] - current_bbox[1]) > 2     # Espace vertical
                                ):
                                    # Finaliser bloc précédent
                                    if current_text.strip():
                                        text_block = TextBlock(
                                            id=f"pdf_block_{page_num}_{block_id}",
                                            text=current_text.strip(),
                                            bbox=BoundingBox(current_bbox[0], current_bbox[1], 
                                                           current_bbox[2], current_bbox[3]),
                                            page=page_num + 1,
                                            font_size=char.get('size', 12.0)
                                        )
                                        text_blocks.append(text_block)
                                        block_id += 1
                                    
                                    # Nouveau bloc
                                    current_text = char['text']
                                    current_bbox = [char['x0'], char['y0'], char['x1'], char['y1']]
                                else:
                                    # Continuer bloc actuel
                                    current_text += char['text']
                                    if current_bbox:
                                        current_bbox[2] = max(current_bbox[2], char['x1'])
                                        current_bbox[3] = max(current_bbox[3], char['y1'])
                                    else:
                                        current_bbox = [char['x0'], char['y0'], char['x1'], char['y1']]
                        
                        # Finaliser dernier bloc
                        if current_text.strip() and current_bbox:
                            text_block = TextBlock(
                                id=f"pdf_block_{page_num}_{block_id}",
                                text=current_text.strip(),
                                bbox=BoundingBox(current_bbox[0], current_bbox[1], 
                                               current_bbox[2], current_bbox[3]),
                                page=page_num + 1
                            )
                            text_blocks.append(text_block)
                        
                        # Texte brut
                        raw_text = page.extract_text() or ""
                        
                        page_data = DocumentPage(
                            page_number=page_num + 1,
                            width=page.width,
                            height=page.height,
                            text_blocks=text_blocks,
                            raw_text=raw_text
                        )
                        
                        pages.append(page_data)
                        
                    except Exception as e:
                        error_msg = f"Erreur page {page_num + 1}: {e}"
                        logger.error(error_msg)
                        errors.append(error_msg)
            
            metadata = {
                'creator': 'pdfplumber',
                'total_text_blocks': sum(len(p.text_blocks) for p in pages)
            }
            
            return ParsedDocument(
                file_path=file_path,
                file_type="PDF",
                pages=pages,
                total_pages=len(pages),
                metadata=metadata,
                parsing_method="pdfplumber",
                parsing_confidence=0.8 if not errors else 0.6,
                errors=errors
            )
            
        except Exception as e:
            logger.error(f"Erreur pdfplumber: {e}")
            return self._empty_document(file_path, "PDF")
    
    def _empty_document(self, file_path: str, file_type: str) -> ParsedDocument:
        """Document vide en cas d'erreur."""
        return ParsedDocument(
            file_path=file_path,
            file_type=file_type,
            pages=[],
            total_pages=0,
            metadata={},
            parsing_method="none",
            parsing_confidence=0.0,
            errors=["Parser non disponible ou échec parsing"]
        )


class DOCXParser:
    """Parser pour fichiers DOCX."""
    
    def __init__(self):
        self.docx_available = DOCX_AVAILABLE
    
    def parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse un fichier DOCX."""
        if not self.docx_available:
            logger.error("python-docx non disponible")
            return self._empty_document(file_path, "DOCX")
        
        try:
            logger.info("📄 Parse DOCX: [FILENAME]")
            
            doc = Document(file_path)
            text_blocks = []
            block_id = 0
            
            # Extraire paragraphes
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # DOCX n'a pas de coordonnées spatiales, simuler
                    bbox = BoundingBox(0, block_id * 20, 500, (block_id + 1) * 20)
                    
                    # Détecter formatage
                    is_bold = any(run.bold for run in para.runs if run.bold)
                    
                    text_block = TextBlock(
                        id=f"docx_block_{block_id}",
                        text=text,
                        bbox=bbox,
                        page=1,  # DOCX traité comme une page
                        is_bold=is_bold
                    )
                    
                    text_blocks.append(text_block)
                    block_id += 1
            
            # Texte brut complet
            raw_text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
            
            page_data = DocumentPage(
                page_number=1,
                width=595,  # A4 approximation
                height=842,
                text_blocks=text_blocks,
                raw_text=raw_text
            )
            
            return ParsedDocument(
                file_path=file_path,
                file_type="DOCX",
                pages=[page_data],
                total_pages=1,
                metadata={'creator': 'python-docx', 'paragraphs': len(doc.paragraphs)},
                parsing_method="python-docx",
                parsing_confidence=0.8,
                errors=[]
            )
            
        except Exception as e:
            logger.error(f"Erreur DOCX: {e}")
            return self._empty_document(file_path, "DOCX")
    
    def _empty_document(self, file_path: str, file_type: str) -> ParsedDocument:
        return ParsedDocument(
            file_path=file_path, file_type=file_type, pages=[], total_pages=0,
            metadata={}, parsing_method="none", parsing_confidence=0.0,
            errors=["Parser DOCX non disponible"]
        )


class UniversalCVParser:
    """Parser universel pour tous formats CV."""
    
    def __init__(self, language: str = 'fr'):
        self.language = language
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
        self.ocr_engine = OCREngine(language)
        
        # Extensions supportées
        self.supported_extensions = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',  # Approximation
            '.odt': 'odt',
            '.png': 'image',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.tiff': 'image',
            '.bmp': 'image'
        }
    
    def parse_document(self, file_path: str, force_ocr: bool = False) -> ParsedDocument:
        """
        Parse un document CV de n'importe quel format supporté.
        
        Args:
            file_path: Chemin vers le fichier
            force_ocr: Forcer OCR même pour PDF natif
            
        Returns:
            ParsedDocument avec contenu extrait
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            logger.error(f"Format non supporté: {extension}")
            return self._empty_document(str(file_path), "UNKNOWN")
        
        file_type = self.supported_extensions[extension]
        logger.info("🔍 Parse document {file_type.upper()}: %s", "[FILENAME]")
        
        if file_type == 'pdf':
            if force_ocr:
                return self._parse_pdf_with_ocr(str(file_path))
            else:
                result = self.pdf_parser.parse_pdf(str(file_path))
                # Si PDF scanné (peu de texte), essayer OCR
                if self._is_scanned_pdf(result):
                    logger.info("PDF scanné détecté, basculement vers OCR")
                    return self._parse_pdf_with_ocr(str(file_path))
                return result
        
        elif file_type == 'docx':
            return self.docx_parser.parse_docx(str(file_path))
        
        elif file_type == 'image':
            return self._parse_image_with_ocr(str(file_path))
        
        elif file_type == 'odt':
            return self._parse_odt(str(file_path))
        
        else:
            return self._empty_document(str(file_path), file_type.upper())
    
    def _is_scanned_pdf(self, parsed_doc: ParsedDocument) -> bool:
        """Détecte si un PDF est scanné (peu de texte extractible)."""
        if not parsed_doc.pages:
            return True
        
        # Heuristique: moins de 100 caractères par page = probablement scanné
        total_chars = sum(len(page.raw_text) for page in parsed_doc.pages)
        avg_chars_per_page = total_chars / len(parsed_doc.pages)
        
        return avg_chars_per_page < 100
    
    def _parse_pdf_with_ocr(self, file_path: str) -> ParsedDocument:
        """Parse PDF avec OCR (pour PDF scannés)."""
        try:
            logger.info("📄 Parse PDF avec OCR: [FILENAME]")
            
            # Convertir PDF en images puis OCR
            import fitz
            
            doc = fitz.open(file_path)
            pages = []
            errors = []
            
            for page_num in range(min(len(doc), 5)):  # Max 5 pages pour performance
                try:
                    page = doc[page_num]
                    
                    # Convertir page en image
                    mat = fitz.Matrix(2, 2)  # 2x zoom pour meilleure qualité
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Sauvegarder temporairement
                    temp_dir = tempfile.mkdtemp()
                    image_path = os.path.join(temp_dir, f"page_{page_num}.png")
                    pix.save(image_path)
                    
                    # OCR sur l'image
                    text_blocks, ocr_confidence = self.ocr_engine.extract_text_with_coordinates(image_path)
                    
                    # Ajuster coordonnées (zoom factor)
                    for block in text_blocks:
                        block.bbox = BoundingBox(
                            block.bbox.left / 2, block.bbox.top / 2,
                            block.bbox.right / 2, block.bbox.bottom / 2
                        )
                        block.page = page_num + 1
                    
                    raw_text = ' '.join(block.text for block in text_blocks)
                    
                    page_data = DocumentPage(
                        page_number=page_num + 1,
                        width=page.rect.width,
                        height=page.rect.height,
                        text_blocks=text_blocks,
                        raw_text=raw_text,
                        image_path=image_path,
                        ocr_confidence=ocr_confidence
                    )
                    
                    pages.append(page_data)
                    
                    # Nettoyage
                    try:
                        os.remove(image_path)
                        if not os.listdir(temp_dir):
                            os.rmdir(temp_dir)
                    except:
                        pass
                    
                except Exception as e:
                    error_msg = f"Erreur OCR page {page_num + 1}: {e}"
                    logger.error(error_msg)
                    errors.append(error_msg)
            
            doc.close()
            
            # Confiance moyenne OCR
            ocr_confidences = [p.ocr_confidence for p in pages if p.ocr_confidence]
            avg_confidence = np.mean(ocr_confidences) if ocr_confidences else 0.5
            
            return ParsedDocument(
                file_path=file_path,
                file_type="PDF",
                pages=pages,
                total_pages=len(pages),
                metadata={
                    'creator': 'OCR',
                    'avg_ocr_confidence': avg_confidence,
                    'total_text_blocks': sum(len(p.text_blocks) for p in pages)
                },
                parsing_method="PDF+OCR",
                parsing_confidence=avg_confidence,
                errors=errors
            )
            
        except Exception as e:
            logger.error(f"Erreur PDF+OCR: {e}")
            return self._empty_document(file_path, "PDF")
    
    def _parse_image_with_ocr(self, file_path: str) -> ParsedDocument:
        """Parse image avec OCR."""
        try:
            logger.info("🖼️ Parse image avec OCR: [FILENAME]")
            
            text_blocks, ocr_confidence = self.ocr_engine.extract_text_with_coordinates(file_path)
            raw_text = ' '.join(block.text for block in text_blocks)
            
            # Dimension image
            with Image.open(file_path) as img:
                width, height = img.size
            
            page_data = DocumentPage(
                page_number=1,
                width=width,
                height=height,
                text_blocks=text_blocks,
                raw_text=raw_text,
                image_path=file_path,
                ocr_confidence=ocr_confidence
            )
            
            return ParsedDocument(
                file_path=file_path,
                file_type="IMAGE",
                pages=[page_data],
                total_pages=1,
                metadata={
                    'creator': 'OCR',
                    'ocr_confidence': ocr_confidence,
                    'image_format': Path(file_path).suffix
                },
                parsing_method="OCR",
                parsing_confidence=ocr_confidence,
                errors=[]
            )
            
        except Exception as e:
            logger.error(f"Erreur OCR image: {e}")
            return self._empty_document(file_path, "IMAGE")
    
    def _parse_odt(self, file_path: str) -> ParsedDocument:
        """Parse ODT (basique)."""
        if not ODT_AVAILABLE:
            logger.error("odfpy non disponible pour ODT")
            return self._empty_document(file_path, "ODT")
        
        try:
            logger.info("📄 Parse ODT: [FILENAME]")
            
            doc = odf_load(file_path)
            paragraphs = doc.getElementsByType(P)
            
            text_blocks = []
            for i, para in enumerate(paragraphs):
                text = str(para).strip()
                if text:
                    bbox = BoundingBox(0, i * 20, 500, (i + 1) * 20)
                    
                    text_block = TextBlock(
                        id=f"odt_block_{i}",
                        text=text,
                        bbox=bbox,
                        page=1
                    )
                    text_blocks.append(text_block)
            
            raw_text = '\n'.join(str(p) for p in paragraphs)
            
            page_data = DocumentPage(
                page_number=1,
                width=595, height=842,
                text_blocks=text_blocks,
                raw_text=raw_text
            )
            
            return ParsedDocument(
                file_path=file_path,
                file_type="ODT",
                pages=[page_data],
                total_pages=1,
                metadata={'creator': 'odfpy'},
                parsing_method="odfpy",
                parsing_confidence=0.7,
                errors=[]
            )
            
        except Exception as e:
            logger.error(f"Erreur ODT: {e}")
            return self._empty_document(file_path, "ODT")
    
    def _empty_document(self, file_path: str, file_type: str) -> ParsedDocument:
        """Document vide."""
        return ParsedDocument(
            file_path=file_path, file_type=file_type, pages=[], total_pages=0,
            metadata={}, parsing_method="none", parsing_confidence=0.0,
            errors=["Parser non disponible ou erreur parsing"]
        )
