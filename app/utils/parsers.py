"""
Document Parsers
===============

Utilitaires pour parser différents types de documents.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional
from loguru import logger
from .text_norm import normalize_text_for_ui


class DocumentParser:
    """Parser pour différents formats de documents."""

    ONGOING_TOKENS = {
        'présent', 'present', 'en cours', 'en-cours', 'current', 'currently',
        'à ce jour', 'a ce jour', 'ce jour', 'today', 'now', 'actuel'
    }

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']

    def parse_document(self, file_path: str) -> str:
        """Parse un document et retourne son contenu texte."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Fichier non trouvé : {file_path}")

        extension = path.suffix.lower()

        if extension == '.txt':
            return self.parse_txt(path)
        elif extension == '.md':
            return self.parse_markdown(path)
        elif extension == '.pdf':
            return self.parse_pdf(path)
        elif extension == '.docx':
            return self.parse_docx(path)
        else:
            raise ValueError(f"Format non supporté : {extension}")

    def _post_process_text(self, text: Optional[str]) -> str:
        """Normalise et nettoie le texte extrait pour usage uniforme."""
        if not text:
            return ''
        normalized = normalize_text_for_ui(text, fix_mojibake=True)
        return normalized.strip()

    def normalize_experience_records(self, experiences: Optional[list]) -> list:
        """Nettoie les expériences pour refléter correctement les statuts "en cours"."""
        if experiences is None:
            return []

        for exp in experiences:
            if not isinstance(exp, dict):
                continue

            end_value = exp.get('end_date')
            current_flag = exp.get('current') is True

            if isinstance(end_value, str):
                lowered = end_value.strip().lower()
                if lowered == '' or any(token in lowered for token in self.ONGOING_TOKENS):
                    current_flag = True
                elif lowered and lowered not in self.ONGOING_TOKENS:
                    current_flag = False
            elif end_value is None:
                current_flag = True

            if current_flag:
                if exp.get('end_date') not in (None, ''):
                    exp['end_date'] = ''
                if exp.get('current') is not True:
                    exp['current'] = True
            else:
                if exp.get('current'):
                    exp['current'] = False
                if end_value is None:
                    exp['end_date'] = ''

        return experiences


    def parse_txt(self, file_path: Path) -> str:
        """Parse un fichier texte."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return self._post_process_text(f.read())
        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            for encoding in ['latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return self._post_process_text(f.read())
                except UnicodeDecodeError:
                    continue
            raise ValueError("Impossible de décoder le fichier texte")

    def parse_markdown(self, file_path: Path) -> str:
        """Parse un fichier Markdown."""
        return self.parse_txt(file_path)

    def parse_pdf(self, file_path: Path) -> str:
        """Parse un fichier PDF avec PyMuPDF si disponible, sinon pypdf."""
        extracted = ''
        try:
            import fitz  # type: ignore
            with fitz.open(file_path) as doc:
                extracted = "\n".join((page.get_text("text") or '').strip() for page in doc)
        except ImportError:
            logger.debug("PyMuPDF non disponible, fallback pypdf")
        except Exception as err:
            logger.warning(f"PyMuPDF a échoué pour {file_path}: {err}")
            extracted = ''

        if not extracted.strip():
            try:
                import pypdf
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    extracted = "\n".join((page.extract_text() or '').strip() for page in reader.pages)
            except ImportError as exc:
                raise ImportError("pypdf n'est pas installé. Installez-le avec: pip install pypdf") from exc
            except Exception as err:
                logger.error(f"Erreur parsing PDF {file_path}: {err}")
                raise ValueError(f"Impossible de lire le PDF : {err}")

        return self._post_process_text(extracted)

    def parse_docx(self, file_path: Path) -> str:
        """Parse un fichier Word (.docx)."""
        try:
            from docx import Document

            doc = Document(file_path)
            chunks = []

            for paragraph in doc.paragraphs:
                chunks.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    cell_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cell_values:
                        chunks.append(' '.join(cell_values))

            raw_text = "\n".join(chunks)
            return self._post_process_text(raw_text)
        except ImportError as exc:
            raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx") from exc
        except Exception as e:
            logger.error(f"Erreur parsing DOCX {file_path}: {e}")
            raise ValueError(f"Impossible de lire le DOCX : {e}")

    def extract_structured_data(self, content: str, document_type: str = "cv") -> Dict[str, Any]:
        """Extrait des données structurées du contenu."""
        if document_type == "cv":
            return self.extract_cv_data(content)
        elif document_type == "offer":
            return self.extract_offer_data(content)
        else:
            return {"content": content}
    
    def extract_cv_data(self, content: str) -> Dict[str, Any]:
        """Extrait les données structurées d'un CV."""
        data = {
            "contact": {},
            "sections": {},
            "skills": [],
            "languages": [],
            "experience": [],
            "education": []
        }
        
        # Extraction email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, content)
        if emails:
            data["contact"]["email"] = emails[0]
        
        # Extraction téléphone (patterns français)
        phone_patterns = [
            r'(?:\+33|0)[1-9](?:[0-9]{8})',
            r'(?:\+33|0)[1-9](?:\s?[0-9]{2}){4}',
            r'(?:\+33|0)[1-9](?:\.[0-9]{2}){4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, content)
            if phones:
                data["contact"]["phone"] = phones[0]
                break
        
        # Extraction LinkedIn
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, content, re.IGNORECASE)
        if linkedin_matches:
            data["contact"]["linkedin"] = linkedin_matches[0]
        
        # Détection de sections par mots-clés
        sections_keywords = {
            "experience": ["expérience", "experience", "emploi", "professionnel", "carrière"],
            "education": ["formation", "education", "études", "diplôme", "université"],
            "skills": ["compétences", "skills", "qualifications", "expertise"],
            "languages": ["langues", "languages", "idiomas"],
            "projects": ["projets", "projects", "réalisations", "portfolio"],
            "certifications": ["certifications", "certificats", "diplômes"]
        }
        
        lines = content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line_clean = line.strip().lower()
            
            # Détection de début de section
            for section, keywords in sections_keywords.items():
                if any(keyword in line_clean for keyword in keywords) and len(line.strip()) < 50:
                    current_section = section
                    data["sections"][section] = []
                    break
            
            # Ajout du contenu à la section courante
            if current_section and line.strip() and not any(keyword in line_clean for keywords_list in sections_keywords.values() for keyword in keywords_list):
                data["sections"][current_section].append(line.strip())
        
        return data
    
    def extract_offer_data(self, content: str) -> Dict[str, Any]:
        """Extrait les données structurées d'une offre d'emploi."""
        data = {
            "title": "",
            "company": "",
            "location": "",
            "contract_type": "",
            "salary": "",
            "requirements": [],
            "responsibilities": [],
            "benefits": []
        }
        
        lines = content.split('\n')
        
        # Essayer d'extraire le titre (généralement dans les premières lignes)
        for line in lines[:5]:
            if line.strip() and len(line.strip()) > 10 and len(line.strip()) < 100:
                # Éviter les lignes trop courtes ou trop longues
                if not data["title"] or len(line.strip()) > len(data["title"]):
                    data["title"] = line.strip()
        
        # Extraction de l'entreprise (recherche de patterns)
        company_patterns = [
            r'(?:entreprise|company|société)\s*:?\s*([^\n]+)',
            r'(?:chez|at)\s+([A-Z][a-zA-Z\s&]+)',
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                data["company"] = matches[0].strip()
                break
        
        # Extraction localisation
        location_keywords = ["paris", "lyon", "marseille", "toulouse", "lille", "bordeaux", "remote", "télétravail"]
        for keyword in location_keywords:
            if keyword in content.lower():
                data["location"] = keyword.title()
                break
        
        # Type de contrat
        contract_types = {
            "CDI": ["cdi", "contrat indéterminé"],
            "CDD": ["cdd", "contrat déterminé"],
            "Stage": ["stage", "stagiaire"],
            "Freelance": ["freelance", "indépendant", "consultant"]
        }
        
        for contract, keywords in contract_types.items():
            if any(keyword in content.lower() for keyword in keywords):
                data["contract_type"] = contract
                break
        
        return data
    
    def clean_text(self, text: str) -> str:
        """Nettoie le texte extrait."""
        # Supprimer les espaces multiples
        text = re.sub(r'\s+', ' ', text)
        
        # Supprimer les lignes vides multiples
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Nettoyer les caractères spéciaux problématiques
        text = text.replace('\x00', '')  # Caractères null
        text = text.replace('\ufeff', '')  # BOM
        
        return text.strip()
    
    def validate_document(self, file_path: str) -> bool:
        """Valide qu'un document peut être parsé."""
        path = Path(file_path)
        
        if not path.exists():
            return False
        
        if path.suffix.lower() not in self.supported_formats:
            return False
        
        # Vérifier la taille (max 10MB)
        if path.stat().st_size > 10 * 1024 * 1024:
            return False
        
        return True
