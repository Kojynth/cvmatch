"""
Loader pour fichiers DOCX
"""

import logging
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
import zipfile

from . import BaseLoader

logger = logging.getLogger(__name__)


class DOCXLoader(BaseLoader):
    """Loader pour fichiers DOCX/DOC"""

    def load(self, file_path: Path) -> Dict[str, Any]:
        """
        Charge un document DOCX et extrait le texte

        Returns:
            Dict avec text, pages, metadata
        """
        logger.debug("📄 Chargement DOCX: %s", "[FILENAME]")

        try:
            # Vérifier si le fichier est un vrai DOCX
            if not self._is_valid_docx(file_path):
                return self._handle_invalid_docx(file_path)

            doc = Document(file_path)

            # Extraire le texte par paragraphes
            paragraphs = []
            full_text = []

            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(
                        {
                            "index": i,
                            "text": text,
                            "style": paragraph.style.name if paragraph.style else None,
                        }
                    )
                    full_text.append(text)

            # Extraire les tableaux
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    tables.append({"index": table_idx, "text": table_text})
                    full_text.append(table_text)

            result = {
                "text": "\n".join(full_text),
                "pages": [
                    {
                        "page_number": 0,
                        "text": "\n".join(full_text),
                        "paragraphs": paragraphs,
                        "tables": tables,
                    }
                ],
                "page_count": 1,
                "needs_ocr": False,
                "format": "docx",
                "metadata": self._extract_docx_metadata(doc, file_path),
            }

            logger.debug(
                f"✅ DOCX chargé: {len(paragraphs)} paragraphes, {len(tables)} tableaux"
            )
            return result

        except Exception as e:
            logger.error(f"❌ Erreur chargement DOCX: {e}")
            return {
                "text": "",
                "pages": [],
                "page_count": 0,
                "needs_ocr": False,
                "format": "docx",
                "metadata": self._extract_metadata(file_path),
                "error": str(e),
            }

    def _is_valid_docx(self, file_path: Path) -> bool:
        """Vérifie si le fichier est un DOCX valide"""
        try:
            with zipfile.ZipFile(file_path, "r") as zip_file:
                return "word/document.xml" in zip_file.namelist()
        except:
            return False

    def _handle_invalid_docx(self, file_path: Path) -> Dict[str, Any]:
        """Gère les fichiers DOCX invalides (souvent des .doc renommés)"""
        logger.warning(
            "⚠️ Fichier DOCX invalide, tentative de lecture basique: %s", "[FILENAME]"
        )

        try:
            # Tentative de lecture comme texte brut
            with open(file_path, "rb") as f:
                content = f.read()
                # Extraire le texte lisible (très basique)
                text = "".join(chr(b) if 32 <= b <= 126 else " " for b in content)
                # Nettoyer le texte
                text = " ".join(text.split())

                return {
                    "text": text,
                    "pages": [{"page_number": 0, "text": text}],
                    "page_count": 1,
                    "needs_ocr": False,
                    "format": "docx_raw",
                    "metadata": self._extract_metadata(file_path),
                    "warnings": ["Fichier DOCX corrompu, extraction basique"],
                }
        except Exception as e:
            logger.error(f"❌ Impossible de lire le fichier: {e}")
            return {
                "text": "",
                "pages": [],
                "page_count": 0,
                "needs_ocr": False,
                "format": "docx",
                "metadata": self._extract_metadata(file_path),
                "error": f"Fichier illisible: {str(e)}",
            }

    def _extract_table_text(self, table) -> str:
        """Extrait le texte d'un tableau DOCX"""
        table_text = []

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                table_text.append(" | ".join(row_text))

        return "\n".join(table_text)

    def _extract_docx_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """Extrait les métadonnées DOCX"""
        metadata = self._extract_metadata(file_path)

        try:
            core_props = doc.core_properties
            metadata.update(
                {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": core_props.created,
                    "modified": core_props.modified,
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": core_props.revision,
                    "language": core_props.language or "",
                }
            )
        except Exception as e:
            logger.warning(f"⚠️ Impossible d'extraire les métadonnées DOCX: {e}")

        return metadata

    def extract_styles(self, file_path: Path) -> Dict[str, Any]:
        """Extrait les informations de style du document"""
        try:
            doc = Document(file_path)
            styles = {}

            for paragraph in doc.paragraphs:
                if paragraph.style and paragraph.text.strip():
                    style_name = paragraph.style.name
                    if style_name not in styles:
                        styles[style_name] = {"count": 0, "examples": []}

                    styles[style_name]["count"] += 1
                    if len(styles[style_name]["examples"]) < 3:
                        styles[style_name]["examples"].append(
                            paragraph.text.strip()[:100]
                        )

            return styles

        except Exception as e:
            logger.error(f"❌ Erreur extraction styles: {e}")
            return {}
